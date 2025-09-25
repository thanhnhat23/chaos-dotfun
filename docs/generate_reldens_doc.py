from docx import Document
from datetime import date

doc = Document()
doc.add_heading('Tài liệu định hướng dự án Reldens', 0)
doc.add_paragraph(f'Ngày tạo: {date.today().strftime("%d/%m/%Y")}')

# 1. Tổng quan
doc.add_heading('1. Tổng quan dự án', level=1)
doc.add_paragraph(
    'Reldens là nền tảng MMORPG mã nguồn mở viết bằng Node.js, kết hợp Colyseus cho đồng bộ thời gian thực, '
    'Phaser 3 cho client 2D và React cho giao diện đăng nhập/chọn nhân vật. Mục tiêu là giúp đội ngũ nhỏ xây dựng '
    'game trực tuyến nhanh chóng và dễ tùy biến.'
)
doc.add_paragraph(
    'Tài liệu này hướng dẫn người mới: cấu trúc thư mục, stack công nghệ, cách chạy dự án và các điểm nên can thiệp '
    'khi muốn mở rộng (ví dụ blockchain).'
)

# 2. Công nghệ
doc.add_heading('2. Công nghệ cốt lõi nên nắm', level=1)
doc.add_paragraph('Danh sách đề xuất học trước khi chỉnh sửa mã nguồn:')
tech_list = [
    'Node.js 18+ và npm – nền tảng chạy server và tool.',
    'Colyseus – thư viện multiplayer, quản lý room và state.',
    'Phaser 3 – engine 2D dựng gameplay trong trình duyệt.',
    'React 18 – giao diện đăng nhập, đăng ký, chọn nhân vật.',
    'Parcel 2 (khuyên dùng 2.13.2) – bundler cho theme front-end.',
    'SCSS + Tailwind (CDN) – lớp style chính của theme.',
    'Hệ sinh thái @reldens (skills, items, storage, utils...) – các module bổ trợ gameplay.',
    'MySQL qua Knex (mặc định) – lưu dữ liệu; có thể đổi sang Postgres/Mongo.',
    'Git, IDE (VSCode), Postman… – công cụ hỗ trợ phát triển.'
]
for item in tech_list:
    doc.add_paragraph(item, style='List Bullet')

# 3. Cấu trúc
doc.add_heading('3. Cấu trúc thư mục & file quan trọng', level=1)
structure = [
    'index.js – điểm khởi chạy, tạo ServerManager.',
    'server.js – export ServerManager để tái sử dụng.',
    'package.json – quản lý dependency và metadata.',
    'bin/ – các lệnh CLI: buildSkeleton, buildCss, buildClient, installer.',
    'dist/ – nơi chứa file build; không chỉnh trực tiếp.',
    'install/ – template sinh .env và hệ thống cài đặt web.',
    'theme/default/index.js – bootstraps React và GameManager.',
    'theme/default/app/ – App.jsx và các trang React (Login, Signup…).',
    'theme/default/css/ – SCSS chia module (base, auth, player-selection...).',
    'theme/default/assets/ – sprite, template HTML, icon, bản đồ.',
    'lib/game/client/ – SceneDynamic, PlayerEngine, RoomEvents…',
    'lib/game/server/ – ServerManager, cấu hình storage, load map.',
    'lib/actions/ – logic kỹ năng và UI hành động.',
    'lib/features/ – module bổ sung như chat, inventory, audio.',
    'lib/users/ – quản lý tài khoản, player profile, đăng nhập.',
    'migrations/ – script tạo bảng (dùng khi không auto generate).'
]
for item in structure:
    doc.add_paragraph(item, style='List Bullet')

# 4. Thiết lập
doc.add_heading('4. Thiết lập môi trường & chạy dự án', level=1)
doc.add_paragraph('Các bước khởi tạo:')
setup_steps = [
    '1. Cài Node.js >= 18.',
    '2. Clone repo và mở trong IDE.',
    '3. Chạy "npm install".',
    '4. Sao chép lib/game/server/install-templates/.env.dist thành .env và điền thông tin RELDENS_DB_*.',
    '5. Chạy "node bin/reldens-commands.js buildSkeleton" để build CSS + client.',
    '6. Chạy "node index.js" để khởi động server.',
    '7. Truy cập http://localhost:8080 để xem giao diện.'
]
for step in setup_steps:
    doc.add_paragraph(step, style='List Number')

doc.add_paragraph('Lệnh hữu ích trong quá trình phát triển:')
commands = [
    'node bin/reldens-commands.js buildCss – build lại SCSS.',
    'node bin/reldens-commands.js buildClient – build lại bundle.',
    'npx parcel@2.13.2 watch theme/default/index.html – xem live khi sửa React/SCSS.',
    'node index.js – chạy server cục bộ.',
    'npm run <script tùy chỉnh> – nếu bạn bổ sung lệnh riêng.'
]
for cmd in commands:
    doc.add_paragraph(cmd, style='List Bullet')

# 5. Luồng hoạt động
doc.add_heading('5. Luồng hoạt động trong mã nguồn', level=1)
doc.add_paragraph(
    'Server: index.js -> ServerManager.start() -> đọc .env -> khởi tạo Express + Colyseus -> load map, features -> lắng nghe kết nối.'
)
doc.add_paragraph(
    'Client: theme/default/index.js dựng React. Khi đăng nhập xong, GameManager client mở SceneDynamic (Phaser) '
    'để điều khiển nhân vật và giao tiếp với Colyseus.'
)
doc.add_paragraph('File quan trọng cần đọc kỹ:')
key_files = [
    'lib/game/client/scene-dynamic.js – xử lý di chuyển, auto target, tương tác.',
    'lib/game/client/room-events.js – quản lý message giữa client và server.',
    'lib/game/server/manager.js – trung tâm cấu hình server.',
    'lib/actions/client/receiver-wrapper.js – nhận message kỹ năng, cập nhật UI.',
    'lib/users/server/login-manager.js – logic xác thực.',
    'lib/features/* – từng module gameplay bổ sung.'
]
for item in key_files:
    doc.add_paragraph(item, style='List Bullet')

# 6. Tùy biến & blockchain
doc.add_heading('6. Tùy biến và tích hợp blockchain', level=1)
doc.add_paragraph(
    'Bạn có toàn quyền chỉnh sửa server Node.js. Khi muốn tích hợp blockchain (NFT item, token phần thưởng, ví người chơi...), hãy tập trung vào:'
)
points = [
    'Luồng đăng nhập: lib/game/server/login-manager.js, lib/users/server/manager.js – nơi có thể gắn xác thực ví.',
    'Chiến đấu, loot, giao dịch: lib/actions/server/*, lib/items – trigger sự kiện on-chain.',
    'Hook server event: ServerManager phát nhiều event (reldens.serverBeforeListen, reldens.serverReady).',
    'Plugin/custom feature: tạo module mới dưới lib/features và đăng ký qua ServerManager.setupCustomServerPlugin.',
    'UI ví: mở rộng theme/default/app/App.jsx, các trang login/signup.',
    'Lưu trữ: lib/game/server/data-server-config.js – thêm bảng lưu hash, địa chỉ ví.',
    'Broadcast: dùng EventsManagerSingleton (từ @reldens/utils) để gửi thông tin tới room hoặc client.'
]
for point in points:
    doc.add_paragraph(point, style='List Bullet')

workflow = [
    '1. Xác định chức năng smart contract (NFT, token, staking...).',
    '2. Viết service/REST layer gọi hợp đồng; có thể đặt trong Express.',
    '3. Bắt sự kiện gameplay -> gửi request on-chain -> nhận phản hồi.',
    '4. Ghi lại kết quả vào DB nội bộ để gameplay không phụ thuộc hoàn toàn blockchain.',
    '5. Cập nhật UI để hiển thị trạng thái giao dịch, số dư, thông báo.'
]
for step in workflow:
    doc.add_paragraph(step, style='List Number')

doc.add_paragraph(
    'Có thể tạo microservice blockchain riêng (ví dụ lib/blockchain/...) rồi import và khởi tạo trong ServerManager.start().' 
    'Cấu trúc module của Reldens hỗ trợ mở rộng khá linh hoạt.'
)

# 7. Tài nguyên
doc.add_heading('7. Tài nguyên tham khảo', level=1)
resources = [
    'Website chính thức: https://www.reldens.com/',
    'Discord cộng đồng: https://discord.gg/HuJMxUY',
    'Repository GitHub: https://github.com/damian-pastorini/reldens',
    'Parcel docs: https://parceljs.org/',
    'Colyseus docs: https://docs.colyseus.io/',
    'Phaser docs: https://phaser.io/'
]
for res in resources:
    doc.add_paragraph(res, style='List Bullet')

doc.add_paragraph(
    'Ghi chú: sau mỗi lần chỉnh SCSS/JS nên build lại; đảm bảo defaultActionKey tồn tại với từng class; kiểm tra .env trước khi deploy.'
)

output_path = 'docs/reldens-project-guide.docx'
doc.save(output_path)
print('Da ghi tai lieu tai {}.'.format(output_path))
